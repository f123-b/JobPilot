from __future__ import annotations

import io
from pathlib import Path


class UnsupportedDocumentError(ValueError):
    pass


def parse_document_bytes(filename: str, content: bytes) -> str:
    suffix = Path(filename or "resume.txt").suffix.lower()
    if suffix in {".txt", ".md", ".text", ".rst"}:
        return content.decode("utf-8", errors="ignore").strip()

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - dependency error is actionable
            raise RuntimeError("pypdf is required for PDF resume parsing") from exc
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required for DOCX resume parsing") from exc
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    paragraphs.append(" | ".join(values))
        return "\n".join(paragraphs).strip()

    raise UnsupportedDocumentError(f"unsupported resume format: {suffix or 'unknown'}")
